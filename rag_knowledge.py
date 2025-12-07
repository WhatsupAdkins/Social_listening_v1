#!/usr/bin/env python3
"""
RAG Knowledge Base for FragPunk Game Comments
Retrieves relevant game domain knowledge to enhance classification accuracy
"""

import os
import json
import pickle
import numpy as np
from typing import List, Dict, Tuple
from docx import Document
from openai import OpenAI
import tiktoken


class FragPunkRAG:
    def __init__(self, 
                 docx_path: str = None,
                 jsonl_paths: List[str] = None,
                 api_key: str = None,
                 embedding_model: str = "text-embedding-3-small",
                 cache_file: str = "fragpunk_rag_cache.pkl"):
        """
        Initialize RAG system for FragPunk domain knowledge
        
        Args:
            docx_path: Path to FragPunk glossary docx file (legacy support)
            jsonl_paths: List of paths to JSONL knowledge files (new format)
            api_key: OpenAI API key
            embedding_model: OpenAI embedding model to use
            cache_file: File to cache embeddings
        """
        self.docx_path = docx_path
        self.jsonl_paths = jsonl_paths or []
        self.api_key = api_key or os.environ.get("OPENAI_API_KEY", "")
        self.embedding_model = embedding_model
        self.cache_file = cache_file
        
        if not self.api_key:
            raise ValueError("OpenAI API key required for RAG system")
        
        if not self.docx_path and not self.jsonl_paths:
            raise ValueError("Either docx_path or jsonl_paths must be provided")
        
        self.client = OpenAI(api_key=self.api_key)
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        
        # Knowledge base storage
        self.chunks = []
        self.embeddings = []
        self.metadata = []
        
    def load_document(self) -> str:
        """Load and extract text from docx file (legacy)"""
        if not self.docx_path:
            return ""
        print(f"📖 Loading FragPunk knowledge from: {self.docx_path}")
        doc = Document(self.docx_path)
        
        # Extract all text
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        print(f"   ✅ Loaded {len(doc.paragraphs)} paragraphs, {len(full_text):,} characters")
        return full_text
    
    def load_jsonl_files(self) -> List[Dict]:
        """Load and parse JSONL knowledge files"""
        if not self.jsonl_paths:
            return []
        
        all_entities = []
        total_files = len(self.jsonl_paths)
        
        print(f"📖 Loading FragPunk knowledge from {total_files} JSONL files...")
        
        for jsonl_path in self.jsonl_paths:
            if not os.path.exists(jsonl_path):
                print(f"   ⚠️  File not found: {jsonl_path}")
                continue
            
            file_name = os.path.basename(jsonl_path)
            entity_count = 0
            
            try:
                with open(jsonl_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        line = line.strip()
                        if not line:
                            continue
                        try:
                            entity = json.loads(line)
                            all_entities.append(entity)
                            entity_count += 1
                        except json.JSONDecodeError as e:
                            print(f"   ⚠️  Error parsing JSON in {file_name}: {e}")
                            continue
                
                print(f"   ✅ Loaded {entity_count} entities from {file_name}")
            except Exception as e:
                print(f"   ⚠️  Error reading {file_name}: {e}")
        
        print(f"   ✅ Total: {len(all_entities)} entities loaded")
        return all_entities
    
    def chunk_text(self, text: str, chunk_size: int = 800, overlap: int = 100) -> List[Dict]:
        """
        Chunk text into overlapping segments
        
        Args:
            text: Full document text
            chunk_size: Target chunk size in tokens
            overlap: Overlap between chunks in tokens
        
        Returns:
            List of chunk dictionaries with text and metadata
        """
        print(f"✂️  Chunking document (chunk_size={chunk_size}, overlap={overlap})...")
        
        # Split by paragraphs first
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        chunks = []
        current_chunk = []
        current_tokens = 0
        chunk_id = 0
        
        for i, para in enumerate(paragraphs):
            para_tokens = len(self.tokenizer.encode(para))
            
            # If paragraph is too long, split it
            if para_tokens > chunk_size:
                # Save current chunk if it has content
                if current_chunk:
                    chunk_text = "\n".join(current_chunk)
                    chunks.append({
                        'id': chunk_id,
                        'text': chunk_text,
                        'tokens': current_tokens,
                        'para_range': (i - len(current_chunk), i)
                    })
                    chunk_id += 1
                    current_chunk = []
                    current_tokens = 0
                
                # Split long paragraph into sentences
                sentences = para.split('. ')
                temp_chunk = []
                temp_tokens = 0
                
                for sent in sentences:
                    sent_tokens = len(self.tokenizer.encode(sent))
                    if temp_tokens + sent_tokens > chunk_size and temp_chunk:
                        chunks.append({
                            'id': chunk_id,
                            'text': '. '.join(temp_chunk) + '.',
                            'tokens': temp_tokens,
                            'para_range': (i, i + 1)
                        })
                        chunk_id += 1
                        temp_chunk = []
                        temp_tokens = 0
                    temp_chunk.append(sent)
                    temp_tokens += sent_tokens
                
                if temp_chunk:
                    chunks.append({
                        'id': chunk_id,
                        'text': '. '.join(temp_chunk) + '.',
                        'tokens': temp_tokens,
                        'para_range': (i, i + 1)
                    })
                    chunk_id += 1
                
                continue
            
            # Check if adding this paragraph exceeds chunk size
            if current_tokens + para_tokens > chunk_size and current_chunk:
                # Save current chunk
                chunk_text = "\n".join(current_chunk)
                chunks.append({
                    'id': chunk_id,
                    'text': chunk_text,
                    'tokens': current_tokens,
                    'para_range': (i - len(current_chunk), i)
                })
                chunk_id += 1
                
                # Start new chunk with overlap
                if overlap > 0 and len(current_chunk) > 1:
                    # Keep last paragraph for overlap
                    current_chunk = [current_chunk[-1], para]
                    current_tokens = len(self.tokenizer.encode(current_chunk[0])) + para_tokens
                else:
                    current_chunk = [para]
                    current_tokens = para_tokens
            else:
                current_chunk.append(para)
                current_tokens += para_tokens
        
        # Don't forget the last chunk
        if current_chunk:
            chunk_text = "\n".join(current_chunk)
            chunks.append({
                'id': chunk_id,
                'text': chunk_text,
                'tokens': current_tokens,
                'para_range': (len(paragraphs) - len(current_chunk), len(paragraphs))
            })
        
        print(f"   ✅ Created {len(chunks)} chunks")
        if len(chunks) > 0:
            total_tokens = sum(c['tokens'] for c in chunks)
            print(f"   📊 Total tokens: {total_tokens:,} (avg: {total_tokens//len(chunks)} per chunk)")
        
        return chunks
    
    def create_embeddings(self, chunks: List[Dict], batch_size: int = 100) -> np.ndarray:
        """
        Create embeddings for all chunks using OpenAI API
        
        Args:
            chunks: List of text chunks
            batch_size: Number of chunks to embed at once
        
        Returns:
            Numpy array of embeddings
        """
        print(f"🔢 Creating embeddings for {len(chunks)} chunks...")
        
        embeddings = []
        
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            texts = [c['text'] for c in batch]
            
            print(f"   Processing batch {i//batch_size + 1}/{(len(chunks)-1)//batch_size + 1}...", end='\r')
            
            # Call OpenAI embedding API
            response = self.client.embeddings.create(
                input=texts,
                model=self.embedding_model
            )
            
            # Extract embeddings
            batch_embeddings = [item.embedding for item in response.data]
            embeddings.extend(batch_embeddings)
        
        print(f"\n   ✅ Created {len(embeddings)} embeddings (dim={len(embeddings[0])})")
        
        return np.array(embeddings)
    
    def build_knowledge_base(self, force_rebuild: bool = False):
        """
        Build the knowledge base from scratch or load from cache
        
        Args:
            force_rebuild: If True, rebuild even if cache exists
        """
        # Check if cache exists
        if not force_rebuild and os.path.exists(self.cache_file):
            print(f"💾 Loading cached knowledge base from {self.cache_file}...")
            with open(self.cache_file, 'rb') as f:
                cache = pickle.load(f)
                self.chunks = cache['chunks']
                self.embeddings = cache['embeddings']
                self.metadata = cache['metadata']
            print(f"   ✅ Loaded {len(self.chunks)} chunks from cache")
            return
        
        # Build from scratch
        print("🏗️  Building knowledge base from scratch...")
        
        # Load from JSONL files if available, otherwise from DOCX
        if self.jsonl_paths:
            entities = self.load_jsonl_files()
            if entities:
                # Convert JSONL entities to chunks
                self.chunks = []
                for i, entity in enumerate(entities):
                    # Use retrieval_text as the main text, fallback to summary, or use name if neither exists
                    text = entity.get('retrieval_text', entity.get('summary', ''))
                    if not text:
                        # If no retrieval_text or summary, use name as the text
                        name = entity.get('name', '')
                        if not name:
                            continue
                        text = name
                    
                    # Build chunk from entity
                    name = entity.get('name', '')
                    entity_type = entity.get('entity_type', '')
                    category = entity.get('category', '')
                    
                    # Build informative chunk text
                    chunk_text = f"{name}"
                    if entity_type:
                        chunk_text += f" ({entity_type})"
                    elif category:
                        # Use category to infer type
                        if 'skin' in category.lower() or 'cosmetic' in category.lower() or 'ornament' in category.lower():
                            chunk_text += f" (cosmetic/skin)"
                        elif 'lancer' in category.lower() or 'character' in category.lower():
                            chunk_text += f" (lancer/character)"
                        elif 'weapon' in category.lower():
                            chunk_text += f" (weapon)"
                        elif 'mode' in category.lower():
                            chunk_text += f" (game mode)"
                    
                    # Add main text
                    if text and text != name:
                        chunk_text += f": {text}"
                    
                    # Add aliases
                    if entity.get('aliases'):
                        aliases_str = ', '.join(entity.get('aliases', []))
                        chunk_text += f" (also known as: {aliases_str})"
                    
                    self.chunks.append({
                        'id': i,
                        'text': chunk_text,
                        'tokens': len(self.tokenizer.encode(chunk_text)),
                        'entity_id': entity.get('id', ''),
                        'entity_type': entity.get('entity_type', ''),
                        'entity_name': entity.get('name', ''),
                        'aliases': entity.get('aliases', []),
                        'metadata': entity,
                        'para_range': (0, 0)  # Not applicable for JSONL
                    })
            else:
                raise ValueError("No entities loaded from JSONL files")
        else:
            # Legacy DOCX loading
            full_text = self.load_document()
            if not full_text:
                raise ValueError("No content loaded from DOCX file")
            # Chunk text
            self.chunks = self.chunk_text(full_text)
        
        if not self.chunks:
            raise ValueError("No knowledge base content loaded. Check your input files.")
        
        # Create embeddings
        embeddings_list = self.create_embeddings(self.chunks)
        self.embeddings = embeddings_list
        
        # Create metadata
        self.metadata = [
            {
                'id': c.get('id', i),
                'tokens': c.get('tokens', 0),
                'entity_id': c.get('entity_id', ''),
                'entity_type': c.get('entity_type', ''),
                'entity_name': c.get('entity_name', ''),
                'aliases': c.get('aliases', []),
                'para_range': c.get('para_range', (0, 0))
            }
            for i, c in enumerate(self.chunks)
        ]
        
        # Save cache
        print(f"💾 Saving knowledge base to {self.cache_file}...")
        with open(self.cache_file, 'wb') as f:
            pickle.dump({
                'chunks': self.chunks,
                'embeddings': self.embeddings,
                'metadata': self.metadata
            }, f)
        print(f"   ✅ Knowledge base built and cached!")
    
    def get_query_embedding(self, query: str) -> np.ndarray:
        """Get embedding for a query string"""
        response = self.client.embeddings.create(
            input=[query],
            model=self.embedding_model
        )
        return np.array(response.data[0].embedding)
    
    def search(self, query: str, top_k: int = 3) -> List[Dict]:
        """
        Search for most relevant chunks given a query
        
        Args:
            query: Search query (comment text)
            top_k: Number of top results to return
        
        Returns:
            List of relevant chunk dictionaries with similarity scores
        """
        if len(self.chunks) == 0:
            raise ValueError("Knowledge base not built. Call build_knowledge_base() first.")
        
        # Get query embedding
        query_embedding = self.get_query_embedding(query)
        
        # Compute cosine similarity
        similarities = np.dot(self.embeddings, query_embedding) / (
            np.linalg.norm(self.embeddings, axis=1) * np.linalg.norm(query_embedding)
        )
        
        # Get top-k indices
        top_indices = np.argsort(similarities)[-top_k:][::-1]
        
        # Return results
        results = []
        for idx in top_indices:
            results.append({
                'chunk_id': self.chunks[idx]['id'],
                'text': self.chunks[idx]['text'],
                'similarity': float(similarities[idx]),
                'metadata': self.metadata[idx]
            })
        
        return results
    
    def get_context_for_comment(self, comment: str, top_k: int = 2, min_similarity: float = 0.3) -> str:
        """
        Get relevant context for a comment
        
        Args:
            comment: Comment text
            top_k: Number of context chunks to retrieve
            min_similarity: Minimum similarity threshold
        
        Returns:
            Formatted context string to inject into prompt
        """
        results = self.search(comment, top_k=top_k)
        
        # Filter by minimum similarity
        relevant_results = [r for r in results if r['similarity'] >= min_similarity]
        
        if not relevant_results:
            return ""
        
        # Format context
        context_parts = []
        for i, result in enumerate(relevant_results, 1):
            context_parts.append(f"[Context {i}]\n{result['text']}")
        
        return "\n\n".join(context_parts)


# Example usage and testing
if __name__ == "__main__":
    import sys
    
    # Test the RAG system
    docx_path = "/Users/zizhengwan/Desktop/FragPunk Glossary and Knowledge Compendium.docx"
    
    if not os.path.exists(docx_path):
        print(f"❌ Document not found: {docx_path}")
        sys.exit(1)
    
    # Initialize RAG
    rag = FragPunkRAG(docx_path=docx_path)
    
    # Build knowledge base
    rag.build_knowledge_base()
    
    # Test queries
    test_comments = [
        "The new skin looks amazing! I'm definitely buying it",
        "Mirage's abilities are too OP, need nerf",
        "The Discipline rifle is so satisfying to use",
        "Love the new battle pass rewards!",
        "Shard cards make the game too random"
    ]
    
    print("\n" + "="*80)
    print("🔍 Testing RAG retrieval with sample comments")
    print("="*80 + "\n")
    
    for comment in test_comments:
        print(f"💬 Comment: {comment}")
        context = rag.get_context_for_comment(comment, top_k=2)
        if context:
            print(f"📚 Retrieved context (truncated):\n{context[:300]}...\n")
        else:
            print("📚 No relevant context found\n")

