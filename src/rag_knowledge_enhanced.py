#!/usr/bin/env python3
"""
Enhanced RAG Knowledge Base with Hybrid Retrieval for FragPunk Game Comments
Fixes entity recall failures by combining semantic search with keyword matching
"""

import os
import json
import pickle
import numpy as np
import re
from typing import List, Dict, Tuple, Set
from collections import defaultdict
from docx import Document
from openai import OpenAI
import tiktoken


class FragPunkRAGEnhanced:
    def __init__(self, 
                 docx_path: str,
                 api_key: str = None,
                 embedding_model: str = "text-embedding-3-small",
                 cache_file: str = "fragpunk_rag_enhanced_cache.pkl"):
        """
        Enhanced RAG system with hybrid retrieval for better entity recall
        
        Args:
            docx_path: Path to FragPunk glossary docx file
            api_key: OpenAI API key
            embedding_model: OpenAI embedding model to use
            cache_file: File to cache embeddings and indices
        """
        self.docx_path = docx_path
        self.api_key = api_key or os.environ.get("OPENAI_API_KEY", "")
        self.embedding_model = embedding_model
        self.cache_file = cache_file
        
        if not self.api_key:
            raise ValueError("OpenAI API key required for RAG system")
        
        self.client = OpenAI(api_key=self.api_key)
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        
        # Knowledge base storage
        self.chunks = []
        self.embeddings = []
        self.metadata = []
        
        # NEW: Keyword index for exact entity matching
        self.entity_index = defaultdict(list)  # entity -> [chunk_ids]
        self.keyword_index = defaultdict(list)  # keyword -> [chunk_ids]
        
        # Known entities (will be populated from document)
        self.entities = set()
    
    def load_document(self) -> str:
        """Load and extract text from docx file"""
        print(f"📖 Loading FragPunk knowledge from: {self.docx_path}")
        doc = Document(self.docx_path)
        
        # Extract all text
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        print(f"   ✅ Loaded {len(doc.paragraphs)} paragraphs, {len(full_text):,} characters")
        return full_text
    
    def extract_entities(self, text: str) -> Set[str]:
        """
        Extract game entities (character names, weapons, abilities, etc.)
        from the document
        """
        print(f"🔍 Extracting game entities...")
        
        entities = set()
        
        # Pattern 1: Capitalized names (likely characters/weapons)
        capitalized = re.findall(r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\b', text)
        
        # Pattern 2: Section headers (often entity names)
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            # Short lines that are capitalized are likely headers/names
            if len(line) < 50 and len(line.split()) <= 4:
                if any(c.isupper() for c in line):
                    entities.add(line.lower())
        
        # Pattern 3: Common game terms
        game_terms = [
            'mirage', 'zephyr', 'dex', 'kismet', 'spider', 'beasty', 'counterfeit',
            'discipline', 'highlife', 'shard', 'lancer', 'ability', 'weapon',
            'assault rifle', 'sniper', 'shotgun', 'pistol', 'smg'
        ]
        
        for term in game_terms:
            if term.lower() in text.lower():
                entities.add(term.lower())
        
        # Add all capitalized words
        for word in capitalized:
            if len(word) > 2:  # Ignore very short words
                entities.add(word.lower())
        
        self.entities = {e for e in entities if len(e) > 2}
        
        print(f"   ✅ Extracted {len(self.entities)} entities")
        if len(self.entities) > 0:
            sample = list(self.entities)[:10]
            print(f"   📝 Sample: {', '.join(sample)}")
        
        return self.entities
    
    def chunk_text(self, text: str, chunk_size: int = 800, overlap: int = 100) -> List[Dict]:
        """Chunk text into overlapping segments"""
        print(f"✂️  Chunking document (chunk_size={chunk_size}, overlap={overlap})...")
        
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        chunks = []
        current_chunk = []
        current_tokens = 0
        chunk_id = 0
        
        for i, para in enumerate(paragraphs):
            para_tokens = len(self.tokenizer.encode(para))
            
            if para_tokens > chunk_size:
                if current_chunk:
                    chunk_text = "\n".join(current_chunk)
                    chunks.append({
                        'id': chunk_id,
                        'text': chunk_text,
                        'tokens': current_tokens,
                        'para_range': (i - len(current_chunk), i)
                    })
                    chunk_id += 1
                    current_chunk = []
                    current_tokens = 0
                
                sentences = para.split('. ')
                temp_chunk = []
                temp_tokens = 0
                
                for sent in sentences:
                    sent_tokens = len(self.tokenizer.encode(sent))
                    if temp_tokens + sent_tokens > chunk_size and temp_chunk:
                        chunks.append({
                            'id': chunk_id,
                            'text': '. '.join(temp_chunk) + '.',
                            'tokens': temp_tokens,
                            'para_range': (i, i + 1)
                        })
                        chunk_id += 1
                        temp_chunk = []
                        temp_tokens = 0
                    temp_chunk.append(sent)
                    temp_tokens += sent_tokens
                
                if temp_chunk:
                    chunks.append({
                        'id': chunk_id,
                        'text': '. '.join(temp_chunk) + '.',
                        'tokens': temp_tokens,
                        'para_range': (i, i + 1)
                    })
                    chunk_id += 1
                
                continue
            
            if current_tokens + para_tokens > chunk_size and current_chunk:
                chunk_text = "\n".join(current_chunk)
                chunks.append({
                    'id': chunk_id,
                    'text': chunk_text,
                    'tokens': current_tokens,
                    'para_range': (i - len(current_chunk), i)
                })
                chunk_id += 1
                
                if overlap > 0 and len(current_chunk) > 1:
                    current_chunk = [current_chunk[-1], para]
                    current_tokens = len(self.tokenizer.encode(current_chunk[0])) + para_tokens
                else:
                    current_chunk = [para]
                    current_tokens = para_tokens
            else:
                current_chunk.append(para)
                current_tokens += para_tokens
        
        if current_chunk:
            chunk_text = "\n".join(current_chunk)
            chunks.append({
                'id': chunk_id,
                'text': chunk_text,
                'tokens': current_tokens,
                'para_range': (len(paragraphs) - len(current_chunk), len(paragraphs))
            })
        
        print(f"   ✅ Created {len(chunks)} chunks")
        total_tokens = sum(c['tokens'] for c in chunks)
        print(f"   📊 Total tokens: {total_tokens:,} (avg: {total_tokens//len(chunks)} per chunk)")
        
        return chunks
    
    def build_keyword_index(self):
        """Build inverted index for keyword matching"""
        print(f"🔨 Building keyword index...")
        
        for chunk in self.chunks:
            chunk_id = chunk['id']
            text_lower = chunk['text'].lower()
            words = re.findall(r'\b\w+\b', text_lower)
            
            # Index all words
            for word in set(words):
                if len(word) > 2:  # Ignore very short words
                    self.keyword_index[word].append(chunk_id)
            
            # Index entities specifically
            for entity in self.entities:
                if entity in text_lower:
                    self.entity_index[entity].append(chunk_id)
        
        print(f"   ✅ Indexed {len(self.keyword_index)} keywords, {len(self.entity_index)} entities")
    
    def create_embeddings(self, chunks: List[Dict], batch_size: int = 100) -> np.ndarray:
        """Create embeddings for all chunks using OpenAI API"""
        print(f"🔢 Creating embeddings for {len(chunks)} chunks...")
        
        embeddings = []
        
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            texts = [c['text'] for c in batch]
            
            print(f"   Processing batch {i//batch_size + 1}/{(len(chunks)-1)//batch_size + 1}...", end='\r')
            
            response = self.client.embeddings.create(
                input=texts,
                model=self.embedding_model
            )
            
            batch_embeddings = [item.embedding for item in response.data]
            embeddings.extend(batch_embeddings)
        
        print(f"\n   ✅ Created {len(embeddings)} embeddings (dim={len(embeddings[0])})")
        
        return np.array(embeddings)
    
    def build_knowledge_base(self, force_rebuild: bool = False):
        """Build the enhanced knowledge base with hybrid indices"""
        if not force_rebuild and os.path.exists(self.cache_file):
            print(f"💾 Loading cached enhanced knowledge base from {self.cache_file}...")
            with open(self.cache_file, 'rb') as f:
                cache = pickle.load(f)
                self.chunks = cache['chunks']
                self.embeddings = cache['embeddings']
                self.metadata = cache['metadata']
                self.entity_index = cache.get('entity_index', defaultdict(list))
                self.keyword_index = cache.get('keyword_index', defaultdict(list))
                self.entities = cache.get('entities', set())
            print(f"   ✅ Loaded {len(self.chunks)} chunks, {len(self.entities)} entities from cache")
            return
        
        print("🏗️  Building enhanced knowledge base from scratch...")
        
        # 1. Load document
        full_text = self.load_document()
        
        # 2. Extract entities
        self.extract_entities(full_text)
        
        # 3. Chunk text
        self.chunks = self.chunk_text(full_text)
        
        # 4. Build keyword index
        self.build_keyword_index()
        
        # 5. Create embeddings
        embeddings_list = self.create_embeddings(self.chunks)
        self.embeddings = embeddings_list
        
        # 6. Create metadata
        self.metadata = [
            {
                'id': c['id'],
                'tokens': c['tokens'],
                'para_range': c['para_range']
            }
            for c in self.chunks
        ]
        
        # 7. Save cache
        print(f"💾 Saving enhanced knowledge base to {self.cache_file}...")
        with open(self.cache_file, 'wb') as f:
            pickle.dump({
                'chunks': self.chunks,
                'embeddings': self.embeddings,
                'metadata': self.metadata,
                'entity_index': dict(self.entity_index),
                'keyword_index': dict(self.keyword_index),
                'entities': self.entities
            }, f)
        print(f"   ✅ Enhanced knowledge base built and cached!")
    
    def get_query_embedding(self, query: str) -> np.ndarray:
        """Get embedding for a query string"""
        response = self.client.embeddings.create(
            input=[query],
            model=self.embedding_model
        )
        return np.array(response.data[0].embedding)
    
    def keyword_search(self, query: str, top_k: int = 5) -> List[int]:
        """
        Perform keyword-based search (BM25-like)
        Returns chunk IDs ranked by keyword relevance
        """
        query_lower = query.lower()
        query_words = set(re.findall(r'\b\w+\b', query_lower))
        
        # Score each chunk
        chunk_scores = defaultdict(float)
        
        # Check for entity matches (highest priority)
        for entity in self.entities:
            if entity in query_lower:
                for chunk_id in self.entity_index.get(entity, []):
                    chunk_scores[chunk_id] += 10.0  # High boost for entity match
        
        # Check for keyword matches
        for word in query_words:
            if word in self.keyword_index:
                for chunk_id in self.keyword_index[word]:
                    chunk_scores[chunk_id] += 1.0
        
        # Sort by score
        sorted_chunks = sorted(chunk_scores.items(), key=lambda x: x[1], reverse=True)
        
        return [chunk_id for chunk_id, score in sorted_chunks[:top_k]]
    
    def hybrid_search(self, query: str, top_k: int = 3, alpha: float = 0.5) -> List[Dict]:
        """
        Hybrid search combining semantic search with keyword matching
        
        Args:
            query: Search query (comment text)
            top_k: Number of top results to return
            alpha: Weight for semantic search (1-alpha for keyword search)
                   0.5 = equal weight, 1.0 = semantic only, 0.0 = keyword only
        
        Returns:
            List of relevant chunk dictionaries with combined scores
        """
        if len(self.chunks) == 0:
            raise ValueError("Knowledge base not built. Call build_knowledge_base() first.")
        
        # 1. Semantic search
        query_embedding = self.get_query_embedding(query)
        semantic_similarities = np.dot(self.embeddings, query_embedding) / (
            np.linalg.norm(self.embeddings, axis=1) * np.linalg.norm(query_embedding)
        )
        
        # Normalize to [0, 1]
        semantic_scores = (semantic_similarities + 1) / 2
        
        # 2. Keyword search
        keyword_chunk_ids = self.keyword_search(query, top_k=len(self.chunks))
        keyword_scores = np.zeros(len(self.chunks))
        for rank, chunk_id in enumerate(keyword_chunk_ids):
            # Decay score by rank
            keyword_scores[chunk_id] = 1.0 / (rank + 1)
        
        # Normalize keyword scores
        if keyword_scores.max() > 0:
            keyword_scores = keyword_scores / keyword_scores.max()
        
        # 3. Combine scores
        combined_scores = alpha * semantic_scores + (1 - alpha) * keyword_scores
        
        # 4. Get top-k indices
        top_indices = np.argsort(combined_scores)[-top_k:][::-1]
        
        # 5. Return results
        results = []
        for idx in top_indices:
            results.append({
                'chunk_id': self.chunks[idx]['id'],
                'text': self.chunks[idx]['text'],
                'similarity': float(combined_scores[idx]),
                'semantic_score': float(semantic_scores[idx]),
                'keyword_score': float(keyword_scores[idx]),
                'metadata': self.metadata[idx]
            })
        
        return results
    
    def search(self, query: str, top_k: int = 3) -> List[Dict]:
        """
        Search using hybrid retrieval (backward compatible interface)
        """
        return self.hybrid_search(query, top_k=top_k, alpha=0.5)
    
    def get_context_for_comment(self, comment: str, top_k: int = 2, min_similarity: float = 0.25) -> str:
        """
        Get relevant context for a comment using hybrid search
        
        Args:
            comment: Comment text
            top_k: Number of context chunks to retrieve
            min_similarity: Minimum combined score threshold (lowered from 0.3)
        
        Returns:
            Formatted context string to inject into prompt
        """
        results = self.search(comment, top_k=top_k)
        
        # Filter by minimum similarity
        relevant_results = [r for r in results if r['similarity'] >= min_similarity]
        
        if not relevant_results:
            return ""
        
        # Format context
        context_parts = []
        for i, result in enumerate(relevant_results, 1):
            context_parts.append(f"[Context {i}]\n{result['text']}")
        
        return "\n\n".join(context_parts)


# Test and comparison
if __name__ == "__main__":
    import sys
    
    docx_path = "/Users/zizhengwan/Desktop/FragPunk Glossary and Knowledge Compendium.docx"
    
    if not os.path.exists(docx_path):
        print(f"❌ Document not found: {docx_path}")
        sys.exit(1)
    
    # Initialize enhanced RAG
    print("="*80)
    print("🚀 Testing Enhanced RAG with Hybrid Retrieval")
    print("="*80 + "\n")
    
    rag = FragPunkRAGEnhanced(docx_path=docx_path)
    rag.build_knowledge_base()
    
    # Test cases with entity mentions
    test_comments = [
        "Zephyr's abilities are too strong",
        "mirage clone is broken",
        "The Discipline rifle needs a buff",
        "Shard cards are too RNG",
        "dex is underpowered"
    ]
    
    print("\n" + "="*80)
    print("🔍 Testing Entity Recall")
    print("="*80 + "\n")
    
    for comment in test_comments:
        print(f"💬 Comment: '{comment}'")
        results = rag.hybrid_search(comment, top_k=2)
        
        print(f"   Found {len(results)} results:")
        for i, r in enumerate(results, 1):
            print(f"   [{i}] Score: {r['similarity']:.3f} (semantic: {r['semantic_score']:.3f}, keyword: {r['keyword_score']:.3f})")
            print(f"       Preview: {r['text'][:100]}...")
        print()

